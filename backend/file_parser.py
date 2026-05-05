import io
import re
import docx
import pandas as pd
import fitz  # PyMuPDF
from fastapi import UploadFile
import easyocr
import numpy as np
import warnings

# Suppress hardware/optimization warnings from Torch for a cleaner demo terminal
warnings.filterwarnings("ignore", category=UserWarning, module="torch.utils.data.dataloader")

# Global reader initialization for efficiency (English language)
# gpu=False ensures it runs on any standard office computer
print("--- Initializing AI OCR Engine ---")
reader = easyocr.Reader(['en'], gpu=False)

async def extract_text_from_upload(file: UploadFile) -> str:
    """
    Handles asynchronous file reading and passes bytes to extraction logic.
    """
    file_bytes = await file.read()
    raw_text = extract_text_from_file(file_bytes, file.filename)
    return clean_extracted_text(raw_text)

def clean_extracted_text(text: str) -> str:
    """
    Optimizes text for AI analysis by removing artifacts and excessive whitespace.
    """
    if not text: return ""
    # Remove non-standard ASCII junk common in PSU scanned documents
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    # Collapse multiple newlines and spaces to save AI tokens
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts text from PDF, Word, and Excel with specialized PDF table handling.
    """
    text = ""
    fn_lower = filename.lower()

    try:
        # 1. ENHANCED PDF HANDLING
        if fn_lower.endswith(".pdf"):
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for i, page in enumerate(doc):
                    # Try high-speed structural text extraction first
                    # 'blocks' preserves paragraph and table-row relationships
                    blocks = page.get_text("blocks")
                    # Sort blocks: Top-to-Bottom, then Left-to-Right
                    blocks.sort(key=lambda b: (b[1], b[0]))
                    extracted = "\n".join([b[4] for b in blocks if b[4].strip()])
                    
                    # Page contains viable text layer
                    if len(extracted.strip()) > 100:
                        text += f"\n--- Page {i+1} ---\n{extracted}\n"
                    else:
                        # OCR FALLBACK: Page is a scanned image (common for BQC pages)
                        # matrix=fitz.Matrix(1.5, 1.5) provides a balance of speed and 
                        # accuracy for reading financial numbers
                        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                        img_data = pix.tobytes("png")
                        
                        # detail=0 returns only text; paragraph=True speeds up logical grouping
                        ocr_results = reader.readtext(img_data, detail=0, paragraph=True)
                        text += f"\n--- Page {i+1} (OCR Scan) ---\n" + "\n".join(ocr_results) + "\n"

        # 2. WORD DOCUMENT HANDLING
        elif fn_lower.endswith((".docx", ".doc")):
            doc_obj = docx.Document(io.BytesIO(file_bytes))
            # Extract from tables (BQC is often in tables in Word docs)
            for table in doc_obj.tables:
                for row in table.rows:
                    text += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
            # Extract standard paragraphs
            text += "\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()])

        # 3. EXCEL HANDLING
        elif fn_lower.endswith((".xlsx", ".xls", ".xlsm")):
            with pd.ExcelFile(io.BytesIO(file_bytes)) as xls:
                for sheet in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet)
                    if not df.empty:
                        text += f"\n[SHEET: {sheet}]\n{df.to_string(index=False)}\n"
        
        else: return "Error: Unsupported file format."

    except Exception as e:
        return f"Error reading file {filename}: {str(e)}"

    return text